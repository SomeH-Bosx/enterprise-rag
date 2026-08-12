"""DOCX / legacy DOC loaders.

Legacy .doc is converted to .docx first (explicit Office/LibreOffice step),
then parsed with python-docx.

Step3.6: tables serialized as Markdown when ENABLE_TABLE_SERIALIZATION=true.
"""

from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document

from src.config.settings import Settings, get_settings
from src.ingestion.loaders._common import docs_or_fail, ensure_file
from src.ingestion.office_convert import prepare_for_load
from src.ingestion.tables import docx_table_to_markdown
from src.services.exceptions import IngestError


def load_docx(path: str | Path, settings: Settings | None = None) -> list[Document]:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:  # pragma: no cover
        raise IngestError("python-docx is required for .docx support") from exc

    cfg = settings or get_settings()
    docx_path = ensure_file(path, expected_suffixes=(".docx",))
    document = DocxDocument(str(docx_path))
    parts: list[str] = []
    table_count = 0

    for para in document.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)

    for table in document.tables:
        if cfg.enable_table_serialization:
            md = docx_table_to_markdown(table)
            if md:
                parts.append(md)
                table_count += 1
        else:
            for row in table.rows:
                cells = [(c.text or "").strip() for c in row.cells]
                line = " | ".join(c for c in cells if c)
                if line:
                    parts.append(line)
            table_count += 1

    text = "\n\n".join(parts).strip()
    docs = [
        Document(
            page_content=text,
            metadata={
                "source": str(docx_path),
                "filename": docx_path.name,
                "file_type": "docx",
                "page": 0,
                "table_count": table_count,
                "ocr_status": "skipped_not_needed",
            },
        )
    ]
    return docs_or_fail(docs, label="DOCX")


def load_doc(path: str | Path, settings: Settings | None = None) -> list[Document]:
    """
    Legacy .doc: explicit convert → .docx, then load_docx.
    Prefer Microsoft Word COM; fallback LibreOffice.
    """
    doc_path = ensure_file(path, expected_suffixes=(".doc",))
    outcome = prepare_for_load(doc_path)
    docs = load_docx(outcome.load_path, settings=settings)
    for d in docs:
        meta = dict(d.metadata or {})
        meta["filename"] = doc_path.name
        meta["source"] = str(doc_path)
        meta["file_type"] = "doc"
        meta["loader"] = f"convert:{outcome.engine}->docx"
        meta["converted"] = outcome.converted
        meta["convert_engine"] = outcome.engine
        d.metadata = meta
    return docs_or_fail(docs, label="DOC")
